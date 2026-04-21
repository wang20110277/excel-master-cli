"""Word (.docx) parser."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from .base import BaseParser, ParseResult, register_parser


class DocxParser(BaseParser):
    """Parse Word .docx files into records."""

    format_name = "docx"

    def parse(self, content: str, source: str = "") -> ParseResult:
        # For docx, content is ignored; we need the file path
        result = ParseResult(source=source, fmt="docx")
        result.warn("Docx parsing requires a file path. Use parse_file() instead.")
        return result

    def parse_file(self, path: str | Path) -> ParseResult:
        """Parse a .docx file."""
        from docx import Document

        p = Path(path)
        result = ParseResult(source=str(p), fmt="docx")
        current_module = ""

        try:
            doc = Document(str(p))
        except Exception as e:
            result.warn(f"Failed to open docx: {e}")
            return result

        # Extract title from first paragraph
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith("Heading"):
                    result.meta["title"] = para.text.strip()
                break

        # Process paragraphs and tables in document order
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                for para in doc.paragraphs:
                    if para._element is element:
                        if para.style.name.startswith("Heading"):
                            current_module = para.text.strip()
                        break

            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        self._parse_table(table, current_module, result)
                        break

        return result

    def _parse_table(self, table, module: str, result: ParseResult) -> None:
        """Extract records from a Word table."""
        from docx import Document

        if len(table.rows) < 2:
            return

        headers = [cell.text.strip() for cell in table.rows[0].cells]

        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            record: dict[str, Any] = {}
            for j, h in enumerate(headers):
                if j < len(values) and values[j]:
                    record[h] = values[j]
            if module:
                record.setdefault("module", module)
                record.setdefault("所属模块", module)
            if record:
                result.add_record(record)


register_parser("docx", DocxParser)
